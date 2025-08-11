from sqlalchemy import create_engine, Column, Integer, String, JSON
from sqlalchemy.orm import declarative_base, sessionmaker
from pgvector.sqlalchemy import Vector
from llama_cpp import Llama
from docx import Document as DocxReader
import PyPDF2
import os

LLM_PATH = "models/Qwen2.5-7B-Instruct-1M-Q3_K_L.gguf"
DATABASE_URL = "postgresql+psycopg2://postgres:rapadura@localhost:5432/postgres"

llm = Llama(model_path=LLM_PATH, embedding=True, n_ctx=2048, n_threads=4)
engine = create_engine(DATABASE_URL)
Session = sessionmaker(bind=engine)
session = Session()
Base = declarative_base()

class Document(Base):
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True)
    filename = Column(String)
    content = Column(String)
    embedding = Column(Vector(3584))
    meta_data = Column(JSON)

Base.metadata.create_all(engine)

def extract_text_from_pdf(filepath):
    text = ""
    with open(filepath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    return text

def extract_text_from_docx(filepath):
    doc = DocxReader(filepath)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

def extract_text_from_txt(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()

def chunk_text(text, max_tokens=300):
    words = text.split()
    return [" ".join(words[i:i + max_tokens]) for i in range(0, len(words), max_tokens)]


def gerar_embedding(texto):
    result = llm.create_embedding(input=texto)
    embedding = result["data"][0]["embedding"]

    if isinstance(embedding[0], list):
        embedding = embedding[0]

    # Garante lista de floats e tamanho correto
    embedding = list(map(float, embedding))
    if len(embedding) != 3584:
        raise ValueError(f"Embedding gerado tem tamanho {len(embedding)}, esperado 3584")
    return embedding


def processar_arquivo(filepath):
    filename = os.path.basename(filepath)

    if filepath.endswith(".pdf"):
        full_text = extract_text_from_pdf(filepath)
    elif filepath.endswith(".docx"):
        full_text = extract_text_from_docx(filepath)
    elif filepath.endswith(".txt"):
        full_text = extract_text_from_txt(filepath)
    else:
        raise ValueError(f"Formato não suportado: {filepath}")

    chunks = chunk_text(full_text)

    for i, chunk in enumerate(chunks):
        emb = gerar_embedding(chunk)
        # Garante lista de floats e tamanho correto
        if not isinstance(emb, list):
            raise ValueError("Embedding não é uma lista")
        if len(emb) != 3584:
            raise ValueError(f"Embedding tem tamanho {len(emb)}, esperado 3584")
        emb = [float(x) for x in emb]
        doc = Document(
            filename=filename,
            content=chunk,
            embedding=emb,
            meta_data={"chunk": i}
        )
        session.add(doc)

    session.commit()
    return len(chunks)

def processar_diretorio(pasta):
    arquivos_processados = {}
    for root, _, files in os.walk(pasta):
        for nome_arquivo in files:
            caminho = os.path.join(root, nome_arquivo)
            try:
                n_chunks = processar_arquivo(caminho)
                arquivos_processados[nome_arquivo] = n_chunks
            except Exception as e:
                arquivos_processados[nome_arquivo] = f"Erro: {str(e)}"
    return arquivos_processados